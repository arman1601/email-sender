import tkinter as tk
from tkinter import filedialog, messagebox
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from openpyxl import load_workbook
from docx import Document
import os

def send_email():
    if not hasattr(send_email, 'excel_file') or not send_email.excel_file:
        messagebox.showerror("Error", "Please select an Excel file!")
        return
    if not hasattr(send_email, 'word_file') or not send_email.word_file:
        messagebox.showerror("Error", "Please select a Word file!")
        return

    from_email = entry_email.get()
    password = entry_password.get()
    smtp_server = entry_smtp.get()
    smtp_port = int(entry_port.get())
    subject = entry_subject.get()

    server = None
    try:
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(from_email, password)

        wb = load_workbook(send_email.excel_file)
        sheet = wb.active

        headers = [cell.value for cell in sheet[1]]  # Получаем заголовки столбцов
        
        for row in sheet.iter_rows(min_row=2, values_only=True):
            data_map = {headers[i]: row[i] for i in range(len(headers)) if row[i] is not None}  # Создаем словарь замен

            if "Email" not in data_map:
                continue
            email = data_map["Email"]

            msg = MIMEMultipart()
            msg['From'] = from_email
            msg['To'] = email
            msg['Subject'] = subject

            # Stex load em anum wordi datan u poxum variablnery
            word_html = ""
            doc = Document(send_email.word_file)
            for paragraph in doc.paragraphs:
                for key, value in data_map.items():
                    paragraph.text = paragraph.text.replace(key, str(value))
            
            for paragraph in doc.paragraphs:
                style = ""
                if paragraph.runs:
                    run = paragraph.runs[0]
                    if run.bold:
                        style += "font-weight: bold; "
                    if run.italic:
                        style += "font-style: italic; "
                    if run.underline:
                        style += "text-decoration: underline; "
                    if run.font.size:
                        style += f"font-size: {run.font.size.pt}pt; "
                
                word_html += f'<p style="{style}">{paragraph.text}</p>'
            
            msg.attach(MIMEText(f"<html><body>{word_html}</body></html>", 'html', 'utf-8'))

            server.sendmail(from_email, email, msg.as_string())

        messagebox.showinfo("Success", "All emails sent successfully!")

    except Exception as e:
        messagebox.showerror("Error", f"Error: {str(e)}")
    finally:
        if server:
            server.quit()

def select_excel_file():
    file = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx")])
    if file:
        send_email.excel_file = file
        label_excel.config(text=f"Selected file: {os.path.basename(file)}")

def select_word_file():
    file = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
    if file:
        send_email.word_file = file
        label_word.config(text=f"Selected file: {os.path.basename(file)}")

root = tk.Tk()
root.title("Email Sender")
root.geometry("600x700")
root.config(bg="#f2f2f2")

header_label = tk.Label(root, text="Email Sender", font=("Arial", 24, "bold"), bg="#f2f2f2", fg="#4a90e2")
header_label.grid(row=0, column=0, columnspan=2, pady=20)

frame_input = tk.Frame(root, bg="#f2f2f2")
frame_input.grid(row=1, column=0, columnspan=2, pady=20, padx=20, sticky="ew")

tk.Label(frame_input, text="Email:", font=("Arial", 12), bg="#f2f2f2").grid(row=1, column=0, padx=10, pady=10, sticky="w")
entry_email = tk.Entry(frame_input, width=40, font=("Arial", 12))
entry_email.grid(row=1, column=1, padx=10, pady=10, sticky="ew")

tk.Label(frame_input, text="Password:", font=("Arial", 12), bg="#f2f2f2").grid(row=2, column=0, padx=10, pady=10, sticky="w")
entry_password = tk.Entry(frame_input, width=40, show="*", font=("Arial", 12))
entry_password.grid(row=2, column=1, padx=10, pady=10, sticky="ew")

tk.Label(frame_input, text="SMTP Server:", font=("Arial", 12), bg="#f2f2f2").grid(row=3, column=0, padx=10, pady=10, sticky="w")
entry_smtp = tk.Entry(frame_input, width=40, font=("Arial", 12))
entry_smtp.grid(row=3, column=1, padx=10, pady=10, sticky="ew")

tk.Label(frame_input, text="Port:", font=("Arial", 12), bg="#f2f2f2").grid(row=4, column=0, padx=10, pady=10, sticky="w")
entry_port = tk.Entry(frame_input, width=40, font=("Arial", 12))
entry_port.grid(row=4, column=1, padx=10, pady=10, sticky="ew")

tk.Label(frame_input, text="Subject:", font=("Arial", 12), bg="#f2f2f2").grid(row=5, column=0, padx=10, pady=10, sticky="w")
entry_subject = tk.Entry(frame_input, width=40, font=("Arial", 12))
entry_subject.grid(row=5, column=1, padx=10, pady=10, sticky="ew")

button_excel = tk.Button(root, text="Select Excel File", command=select_excel_file, font=("Arial", 12), bg="#4a90e2", fg="white", relief="raised", width=20)
button_excel.grid(row=6, column=0, pady=10, padx=20, sticky="ew")
label_excel = tk.Label(root, text="No file selected", font=("Arial", 12), bg="#f2f2f2")
label_excel.grid(row=6, column=1, pady=10, padx=20, sticky="ew")

button_word = tk.Button(root, text="Select Word File", command=select_word_file, font=("Arial", 12), bg="#4a90e2", fg="white", relief="raised", width=20)
button_word.grid(row=7, column=0, pady=10, padx=20, sticky="ew")
label_word = tk.Label(root, text="No file selected", font=("Arial", 12), bg="#f2f2f2")
label_word.grid(row=7, column=1, pady=10, padx=20, sticky="ew")

button_send = tk.Button(root, text="Send Emails", command=send_email, font=("Arial", 12), bg="#4a90e2", fg="white", relief="raised", width=20)
button_send.grid(row=8, column=0, columnspan=2, pady=20, padx=20, sticky="ew")

root.grid_rowconfigure(1, weight=1)
root.grid_columnconfigure(1, weight=1)

root.mainloop()
